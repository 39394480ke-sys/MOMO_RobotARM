from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = "/Users/ke/Library/Mobile Documents/com~apple~CloudDocs/Code/机械臂/MoMo答辩演讲稿_4分钟.docx"

sections = [
    ("开场与问题", "0:00-0:35", "P1-P7", """各位评委老师大家好，我们是来自深圳大学的MOMO队。今天汇报的项目是MoMo边缘智能摄影机械臂交互控制系统。
在电商商拍、短视频创作和摄影教学中，横移、环绕、跟随等镜头需要反复调参；同时，画面、轨迹、设备状态和运行日志往往无法统一记录。现有跟拍云台便携但复用能力弱，智能滑轨自由度有限，专业MoCo设备又价格高、部署复杂。因此，我们希望将一套可感知、可联网、可记录、可复用的自动摄影能力，下沉到桌面场景。"""),
    ("产品与架构", "0:35-1:25", "P8-P13", """MoMo的硬件主体是五轴机械臂加直线导轨，以SC171V3 AIoT开发套件作为边缘中枢，接入USB摄像头、串口总线舅机和手机或相机支架。
系统形成了“感知、决策、执行、回传”的闭环：摄像头获取主体位置，边缘端负责视觉处理、任务调度和安全校验，机械臂与导轨执行运镜，Web、GUI和手机端则同步画面、进度、状态与日志。云端AI只负责理解用户意图，真实硬件始终由本地工具链控制。"""),
    ("核心功能演示", "1:25-2:35", "P14-P18", """下面介绍四项核心功能。
第一，轨迹录制与复用。我们在手动控制或示教过程中，记录关节角度、动作顺序和时间参数，生成JSON动作序列。下次拍摄同类商品时，可以一键调用，并支持暂停、继续和停止。
第二，AI对话运镜。用户只需输入“环绕拍摄商品”等自然语言，云端模型识别动作类型与参数，边缘端将其映射为预定义能力，再经dry-run、限位、速度、急停和人工确认后执行。
第三，视觉跟随。系统使用YuNet人脸检测或手动框选锁定目标，计算主体与画面中心的偏差，过滤微小抖动后转换为调整指令。
第四，多端协同。同一局域网内，不同终端通过WebSocket共享任务进度和设备状态，便于远程查看与协同操作。"""),
    ("特色创新", "2:35-3:35", "P19-P23", """我们的创新不是单独增加一个AI按钮，而是构建摄影数据的完整闭环。
首先，执行过程同步采集画面、舅机状态、导轨位置和运行日志，让拍摄从“靠经验操作”变为可追踪、可复盘的数据链路。
其次，系统实现云边安全隔离：云端负责“想怎么拍”，边缘端决定“能不能动、怎么安全动”，避免大模型直接写入真实舅机。
最后，动作轨迹、标定参数和日志被沉淀为动作库与模板。完成一次示教后，同类任务可直接复用，并根据执行反馈继续优化，形成“采集、执行、反馈、沉淀、再调用”的循环。"""),
    ("总结展望", "3:35-4:00", "P24-P27", """MoMo面向电商商拍、短视频创作、摄影教学和AIoT实训场景。它不是替代专业影视机器人，而是以更轻量、更低成本的方式，提供可联网、可复用、可智能控制的桌面级自动摄影能力。
下一阶段，我们将增加自动补光、丰富运镜模板和视觉策略，并逐步走向多机位协同与云端模板市场。我们的汇报到此结束，感谢各位评委老师的聆听！"""),
]


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def style_run(run, size=11, bold=False, color="222222"):
    run.font.name = "Arial Unicode MS"
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Arial Unicode MS")
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


doc = Document()
sec = doc.sections[0]
sec.top_margin = Inches(0.75)
sec.bottom_margin = Inches(0.7)
sec.left_margin = Inches(0.9)
sec.right_margin = Inches(0.9)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Arial Unicode MS"
normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial Unicode MS")
normal.font.size = Pt(11)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.25

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.paragraph_format.space_after = Pt(5)
style_run(title.add_run("MoMo 答辩演讲稿"), 22, True, "1679A7")
sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.paragraph_format.space_after = Pt(12)
style_run(sub.add_run("边缘智能摄影机械臂交互控制系统 | 目标时长：4分钟"), 11, False, "5B6573")

table = doc.add_table(rows=1, cols=3)
table.autofit = False
widths = [Inches(1.4), Inches(1.5), Inches(3.6)]
labels = ["建议语速", "总时长", "使用方式"]
vals = ["230-250字/分钟", "3分50秒-4分10秒", "换页提示只用于排练，不念出"]
for i, cell in enumerate(table.rows[0].cells):
    cell.width = widths[i]
    set_cell_shading(cell, "E8F3F8")
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    style_run(p.add_run(labels[i] + "\n"), 9, True, "1679A7")
    style_run(p.add_run(vals[i]), 9.5, False, "222222")

for idx, (heading, timing, slides, body) in enumerate(sections):
    h = doc.add_paragraph()
    h.paragraph_format.space_before = Pt(13 if idx else 15)
    h.paragraph_format.space_after = Pt(5)
    style_run(h.add_run(heading), 15, True, "1679A7")
    style_run(h.add_run(f"   {timing}  |  {slides}"), 9.5, False, "6E7781")
    for para_text in [x.strip() for x in body.strip().split("\n") if x.strip()]:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(5)
        p.paragraph_format.line_spacing = 1.25
        style_run(p.add_run(para_text), 11, False, "222222")

note = doc.add_paragraph()
note.paragraph_format.space_before = Pt(12)
note.paragraph_format.space_after = Pt(0)
style_run(note.add_run("排练提示："), 10, True, "C04C36")
style_run(note.add_run("功能演示时只指出动作结果，不逐项描述界面；若现场演示超过20秒，可删去“视觉跟随”段落中的算法细节。"), 10, False, "444444")

header = sec.header.paragraphs[0]
header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
style_run(header.add_run("2026全国大学生物联网设计竞赛 | MOMO队"), 8.5, False, "76808A")
footer = sec.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
style_run(footer.add_run("MoMo 答辩演讲稿"), 8.5, False, "76808A")

doc.save(OUT)
print(OUT)
